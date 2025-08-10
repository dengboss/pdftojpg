import os
import sys
from PyQt5.QtWidgets import (QApplication, QMainWindow, QVBoxLayout, QHBoxLayout, 
                             QWidget, QPushButton, QFileDialog, QLabel, QSpinBox,
                             QComboBox, QProgressBar, QMessageBox, QTextEdit, 
                             QListWidget, QListWidgetItem, QAbstractItemView)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QDragEnterEvent, QDropEvent
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import io
import shutil
try:
    import win32com.client
except ImportError:
    win32com = None


class ConverterThread(QThread):
    progress_updated = pyqtSignal(int, str)
    file_started = pyqtSignal(str)
    conversion_finished = pyqtSignal(bool, str, int, int)
    
    def __init__(self, file_list, output_dir, format_type, pages, dpi=300):
        super().__init__()
        self.file_list = file_list
        self.output_dir = output_dir
        self.format_type = format_type
        self.pages = pages
        self.dpi = dpi
        
        # 确保输出目录存在
        if not os.path.exists(self.output_dir):
            try:
                os.makedirs(self.output_dir)
            except Exception as e:
                raise Exception(f"无法创建输出目录: {str(e)}")
                
        # 检查目录写入权限
        if not os.access(self.output_dir, os.W_OK):
            raise Exception(f"输出目录无写入权限: {self.output_dir}")
        
    def run(self):
        total_files = len(self.file_list)
        completed_files = 0
        failed_files = 0
        
        for file_path in self.file_list:
            try:
                self.file_started.emit(os.path.basename(file_path))
                
                # 检查文件是否存在
                if not os.path.exists(file_path):
                    raise FileNotFoundError(f"文件不存在: {file_path}")
                
                # 检查文件大小
                if os.path.getsize(file_path) == 0:
                    raise ValueError("文件为空")
                
                # 根据文件扩展名和内容类型选择转换方法
                file_ext = os.path.splitext(file_path)[1].lower()
                
                if file_ext == '.pdf':
                    self.convert_pdf(file_path)
                elif file_ext in ['.docx', '.doc']:
                    try:
                        import zipfile
                        import mimetypes
                        
                        # 对于.doc文件，自动调用Word转换为docx
                        if file_ext == '.doc':
                            docx_path = None
                            try:
                                if win32com is not None:
                                    word = win32com.client.Dispatch('Word.Application')
                                    word.Visible = False
                                    docx_path = file_path + "_tmp_autoconvert.docx"
                                    doc = word.Documents.Open(file_path)
                                    doc.SaveAs(docx_path, FileFormat=16)  # 16 = wdFormatDocumentDefault (docx)
                                    doc.Close()
                                    word.Quit()
                                    # 用docx逻辑处理
                                    self.convert_word(docx_path)
                                    # 删除临时文件
                                    if os.path.exists(docx_path):
                                        os.remove(docx_path)
                                else:
                                    raise ImportError('未安装pywin32，无法自动调用Word转换doc为docx')
                            except Exception as e:
                                if docx_path and os.path.exists(docx_path):
                                    os.remove(docx_path)
                                raise ValueError(f"自动调用Word转换doc为docx失败: {str(e)}。请用Word手动另存为docx后再试。")
                        
                        # 对于.docx文件，进行详细检查
                        elif file_ext == '.docx':
                            try:
                                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                                    file_list = zip_ref.namelist()
                                    # 检查是否包含Word文档的核心文件
                                    if 'word/document.xml' not in file_list:
                                        # 检查是否是主题文件
                                        if any('theme' in f.lower() for f in file_list):
                                            raise ValueError("文件格式错误: 这是一个Office主题文件，不是Word文档")
                                        else:
                                            raise ValueError("文件格式错误: 这不是有效的Word文档")
                            except zipfile.BadZipFile:
                                raise ValueError("文件格式错误: 文件已损坏或不是有效的Word文档")
                            
                            # 尝试用python-docx打开
                            try:
                                doc = Document(file_path)
                                # 检查文档是否有内容
                                has_content = (len(doc.paragraphs) > 0 and any(p.text.strip() for p in doc.paragraphs)) or len(doc.tables) > 0
                                if not has_content:
                                    raise ValueError("Word文档为空或无有效内容")
                                self.convert_word(file_path)
                            except Exception as e:
                                if 'themeManager' in str(e) or 'theme' in str(e):
                                    raise ValueError("文件格式错误: 这是一个Office主题文件，不是Word文档")
                                elif 'package' in str(e).lower():
                                    raise ValueError("文件格式错误: 文件包结构异常，可能已损坏")
                                else:
                                    raise ValueError(f"Word文件处理失败: {str(e)}")
                                    
                    except ValueError as e:
                        raise e
                    except Exception as e:
                        raise ValueError(f"Word文件处理异常: {str(e)}")

                else:
                    raise ValueError(f"不支持的文件格式: {file_ext}\n当前仅支持PDF和Word文件格式")
                    
                completed_files += 1
                
            except FileNotFoundError as e:
                error_msg = f"文件未找到: {str(e)}"
                self.conversion_finished.emit(False, error_msg, 0, 1)
                failed_files += 1
                
            except ValueError as e:
                error_msg = f"文件格式问题: {os.path.basename(file_path)} - {str(e)}"
                self.conversion_finished.emit(False, error_msg, 0, 1)
                failed_files += 1
                
            except Exception as e:
                error_msg = f"转换失败: {os.path.basename(file_path)} - {str(e)}"
                self.conversion_finished.emit(False, error_msg, 0, 1)
                failed_files += 1
        
        total_processed = completed_files + failed_files
        self.conversion_finished.emit(
            failed_files == 0, 
            f"批量转换完成: 成功{completed_files}个, 失败{failed_files}个", 
            completed_files, 
            total_files
        )
    
    def convert_pdf(self, file_path):
        try:
            doc = fitz.open(file_path)
            if len(doc) == 0:
                raise ValueError("PDF文件为空或无法读取")
                
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            total_pages = len(doc) if self.pages == "all" else min(self.pages, len(doc))
            pages_to_convert = range(total_pages) if self.pages == "all" else range(min(self.pages, len(doc)))
            
            for i, page_num in enumerate(pages_to_convert):
                page = doc[page_num]
                mat = fitz.Matrix(self.dpi/72, self.dpi/72)
                pix = page.get_pixmap(matrix=mat)
                
                output_path = os.path.join(self.output_dir, 
                                         f"{base_name}_page_{page_num+1}.{self.format_type}")
                pix.save(output_path)
                
                progress = int((i + 1) / len(pages_to_convert) * 100)
                self.progress_updated.emit(progress, f"{base_name} - 第{page_num+1}页")
                
            doc.close()
        except Exception as e:
            raise Exception(f"PDF转换错误: {str(e)}")
    
    def create_text_image(self, text, base_name):
        """创建包含文本的图片"""
        from PIL import Image, ImageDraw, ImageFont
        
        # 设置图片参数
        width = 800
        height = 600
        background_color = 'white'
        text_color = 'black'
        
        # 创建图片
        img = Image.new('RGB', (width, height), background_color)
        draw = ImageDraw.Draw(img)
        
        # 尝试使用系统字体，如果失败则使用默认字体
        try:
            font = ImageFont.truetype('arial.ttf', 20)
        except:
            try:
                font = ImageFont.truetype('simhei.ttf', 20)
            except:
                font = ImageFont.load_default()
        
        # 计算文本位置（居中显示）
        lines = []
        if len(text) > 50:
            # 长文本分行
            words = text.split()
            line = ""
            for word in words:
                if len(line + word) < 50:
                    line += word + " "
                else:
                    lines.append(line.strip())
                    line = word + " "
            if line:
                lines.append(line.strip())
        else:
            lines = [text]
        
        # 绘制文本
        y_position = 100
        for line in lines:
            draw.text((50, y_position), line, fill=text_color, font=font)
            y_position += 30
        
        return img
    
    def create_text_document_image(self, text_content, base_name):
        """创建Word文档文本内容的图片"""
        from PIL import Image, ImageDraw, ImageFont
        
        # 设置图片参数
        width = 800
        min_height = 600
        background_color = 'white'
        text_color = 'black'
        
        # 计算所需高度
        lines = text_content
        total_lines = len(lines)
        height = max(min_height, total_lines * 25 + 100)
        
        # 创建图片
        img = Image.new('RGB', (width, height), background_color)
        draw = ImageDraw.Draw(img)
        
        # 尝试使用系统字体
        try:
            font = ImageFont.truetype('arial.ttf', 16)
        except:
            try:
                font = ImageFont.truetype('simhei.ttf', 16)
            except:
                font = ImageFont.load_default()
        
        # 绘制标题
        title_font = ImageFont.load_default()
        try:
            title_font = ImageFont.truetype('arial.ttf', 20)
        except:
            try:
                title_font = ImageFont.truetype('simhei.ttf', 20)
            except:
                pass
        
        draw.text((50, 30), f"Word文档: {base_name}", fill=text_color, font=title_font)
        
        # 绘制文本内容
        y_position = 80
        for line in lines:
            if y_position < height - 30:  # 确保不超出图片边界
                draw.text((50, y_position), line, fill=text_color, font=font)
                y_position += 25
        
        return img

    def convert_word(self, file_path):
        try:
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            
            # 方案1：使用Word转换为PDF，保持完整格式
            if win32com is not None:
                try:
                    print(f"尝试使用Word转换: {file_path}")
                    
                    # 创建临时PDF文件
                    temp_pdf_path = os.path.join(self.output_dir, f"{base_name}_temp.pdf")
                    
                    # 使用Word打开文档并另存为PDF
                    word = win32com.client.Dispatch('Word.Application')
                    word.Visible = False
                    
                    # 获取绝对路径
                    abs_file_path = os.path.abspath(file_path)
                    print(f"Word打开文件: {abs_file_path}")
                    
                    doc = word.Documents.Open(abs_file_path)
                    print(f"Word文件打开成功，开始保存为PDF")
                    
                    doc.SaveAs(temp_pdf_path, FileFormat=17)  # 17 = wdFormatPDF
                    print(f"PDF保存成功: {temp_pdf_path}")
                    doc.Close()
                    word.Quit()
                    
                    # 检查PDF文件是否存在
                    if not os.path.exists(temp_pdf_path):
                        raise Exception("PDF文件未生成")
                    
                    print(f"开始处理PDF文件: {temp_pdf_path}")
                    
                    # 使用PyMuPDF处理PDF
                    pdf_doc = fitz.open(temp_pdf_path)
                    total_pages = len(pdf_doc)
                    print(f"PDF总页数: {total_pages}")
                    
                    # 确定要处理的页数
                    pages_to_process = total_pages if self.pages == "all" else min(self.pages, total_pages)
                    
                    for page_num in range(pages_to_process):
                        page = pdf_doc.load_page(page_num)
                        
                        # 设置缩放比例以获得高质量图片
                        mat = fitz.Matrix(self.dpi/72, self.dpi/72)
                        pix = page.get_pixmap(matrix=mat)
                        
                        # 转换为PIL图片
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))
                        
                        # 保存图片
                        output_path = os.path.join(self.output_dir, f"{base_name}_page_{page_num+1}.{self.format_type}")
                        img.save(output_path, format=('JPEG' if self.format_type == 'jpeg' else 'PNG'))
                        
                        progress = int((page_num + 1) / pages_to_process * 100)
                        self.progress_updated.emit(progress, f"{base_name} - 页面 {page_num+1}")
                    
                    pdf_doc.close()
                    
                    # 删除临时PDF文件
                    if os.path.exists(temp_pdf_path):
                        os.remove(temp_pdf_path)
                    
                    print(f"Word转换完成，共处理 {pages_to_process} 页")
                    return
                    
                except Exception as e:
                    print(f"Word转PDF失败，回退到文本模式: {e}")
                    # 如果Word转换失败，回退到原来的文本模式
                    pass
            
            # 方案2：回退到原来的文本模式（保持原有功能作为备选）
            doc = Document(file_path)
            if not doc.paragraphs and not doc.tables:
                raise ValueError("Word文档为空或无法读取内容")
            
            # 提取所有内嵌图片
            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        image = Image.open(io.BytesIO(image_data))
                        images.append(image)
                    except Exception as e:
                        print(f"提取图片失败: {e}")
                        continue
            
            # 创建文档文本内容的图片表示，增强编码处理
            text_content = []
            for para in doc.paragraphs:
                try:
                    text = para.text.strip()
                    if text:
                        # 处理文本编码
                        if not isinstance(text, str):
                            text = str(text)
                        text = text.encode('utf-8', errors='ignore').decode('utf-8')
                        text_content.append(text)
                except Exception as e:
                    print(f"处理段落文本失败: {e}")
                    continue
            
            # 为表格内容添加文本，增强编码处理
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        try:
                            text = cell.text.strip()
                            if text:
                                # 处理文本编码
                                if not isinstance(text, str):
                                    text = str(text)
                                text = text.encode('utf-8', errors='ignore').decode('utf-8')
                                row_text.append(text)
                        except Exception as e:
                            print(f"处理表格单元格文本失败: {e}")
                            continue
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            # 创建文本内容的图片
            if text_content:
                text_img = self.create_text_document_image(text_content, base_name)
                images.insert(0, text_img)
            
            # 如果没有内容，创建提示图片
            if not images:
                img = self.create_text_image(f"Word文档: {base_name} (无内容)")
                images.append(img)
            
            total_images = len(images) if self.pages == "all" else min(self.pages, len(images))
            images_to_save = images[:total_images]
            
            for i, img in enumerate(images_to_save):
                output_path = os.path.join(self.output_dir, f"{base_name}_content_{i+1}.{self.format_type}")
                img.save(output_path, format=('JPEG' if self.format_type == 'jpeg' else 'PNG'))
                
                progress = int((i + 1) / len(images_to_save) * 100)
                self.progress_updated.emit(progress, f"{base_name} - 图片 {i+1}")
                
        except Exception as e:
            raise Exception(f"Word转换错误: {str(e)}")


class DocumentConverter(QMainWindow):
    def __init__(self):
        super().__init__()
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle('文档批量转图片工具')
        self.setGeometry(100, 100, 800, 600)
        
        # 主窗口部件
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        
        # 布局
        layout = QVBoxLayout()
        
        # 文件选择区域
        file_layout = QHBoxLayout()
        self.file_label = QLabel("选择文件:")
        self.select_btn = QPushButton("添加文件")
        self.select_btn.clicked.connect(self.select_files)
        self.select_folder_btn = QPushButton("添加文件夹")
        self.select_folder_btn.clicked.connect(self.select_folder)
        self.clear_btn = QPushButton("清空列表")
        self.clear_btn.clicked.connect(self.clear_files)
        
        file_layout.addWidget(self.file_label)
        file_layout.addWidget(self.select_btn)
        file_layout.addWidget(self.select_folder_btn)
        file_layout.addWidget(self.clear_btn)
        
        # 输出目录
        output_layout = QHBoxLayout()
        self.output_label = QLabel("输出目录:")
        self.output_path_label = QLabel("未选择目录")
        self.output_btn = QPushButton("选择目录")
        self.output_btn.clicked.connect(self.select_output_dir)
        
        output_layout.addWidget(self.output_label)
        output_layout.addWidget(self.output_path_label, 1)
        output_layout.addWidget(self.output_btn)
        
        # 格式选择
        format_layout = QHBoxLayout()
        format_layout.addWidget(QLabel("输出格式:"))
        self.format_combo = QComboBox()
        self.format_combo.addItems(['jpeg', 'png'])
        format_layout.addWidget(self.format_combo)
        
        # 页数选择
        pages_layout = QHBoxLayout()
        pages_layout.addWidget(QLabel("页数:"))
        self.pages_spin = QSpinBox()
        self.pages_spin.setMinimum(1)
        self.pages_spin.setMaximum(999)
        self.pages_spin.setValue(1)
        pages_layout.addWidget(self.pages_spin)
        
        self.all_pages_btn = QPushButton("全部")
        self.all_pages_btn.clicked.connect(lambda: self.pages_spin.setValue(999))
        pages_layout.addWidget(self.all_pages_btn)
        
        # 清晰度选择
        dpi_layout = QHBoxLayout()
        dpi_layout.addWidget(QLabel("清晰度:"))
        self.dpi_combo = QComboBox()
        self.dpi_combo.addItems(['低', '中', '高'])
        self.dpi_combo.setCurrentText('中')
        dpi_layout.addWidget(self.dpi_combo)
        
        # 文件列表
        self.file_list = QListWidget()
        self.file_list.setSelectionMode(QAbstractItemView.ExtendedSelection)
        
        # 转换按钮
        self.convert_btn = QPushButton("开始转换")
        self.convert_btn.clicked.connect(self.start_conversion)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setTextVisible(True)
        
        # 日志区域
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        
        # 移除选中按钮
        remove_layout = QHBoxLayout()
        self.remove_btn = QPushButton("移除选中")
        self.remove_btn.clicked.connect(self.remove_selected_files)
        remove_layout.addWidget(self.remove_btn)
        remove_layout.addStretch()
        
        # 添加所有组件到布局
        layout.addLayout(file_layout)
        layout.addWidget(QLabel("文件列表:"))
        layout.addWidget(self.file_list)
        layout.addLayout(remove_layout)
        layout.addLayout(output_layout)
        layout.addLayout(format_layout)
        layout.addLayout(pages_layout)
        layout.addLayout(dpi_layout)
        layout.addWidget(self.convert_btn)
        layout.addWidget(self.progress_bar)
        layout.addWidget(QLabel("日志:"))
        layout.addWidget(self.log_text)
        
        main_widget.setLayout(layout)
        
        # 设置拖放支持
        self.setAcceptDrops(True)
        
        # 初始化变量
        self.output_dir = ""
        self.converter_thread = None
        
    def select_files(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, "选择文件", "", 
            "PDF文件 (*.pdf);;Word文件 (*.docx);;所有文件 (*.*)"
        )
        for file_path in files:
            if file_path.lower().endswith('.doc'):
                QMessageBox.warning(self, "格式不支持", 
                    f"文件 {os.path.basename(file_path)} 是旧版.doc格式，\n"
                    f"请先将文件转换为.docx格式后再导入。")
                continue
                
            if not self.is_file_in_list(file_path):
                item = QListWidgetItem(file_path)
                item.setData(Qt.UserRole, file_path)
                self.file_list.addItem(item)
    
    def select_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, "选择文件夹")
        if folder_path:
            for root, dirs, files in os.walk(folder_path):
                for file in files:
                    if file.lower().endswith(('.pdf', '.docx')):
                        file_path = os.path.join(root, file)
                        if not self.is_file_in_list(file_path):
                            item = QListWidgetItem(file_path)
                            item.setData(Qt.UserRole, file_path)
                            self.file_list.addItem(item)
    
    def is_file_in_list(self, file_path):
        for i in range(self.file_list.count()):
            if self.file_list.item(i).data(Qt.UserRole) == file_path:
                return True
        return False
    
    def clear_files(self):
        self.file_list.clear()
    
    def remove_selected_files(self):
        for item in self.file_list.selectedItems():
            self.file_list.takeItem(self.file_list.row(item))
    
    def select_output_dir(self):
        dir_path = QFileDialog.getExistingDirectory(self, "选择输出目录")
        if dir_path:
            self.output_dir = dir_path
            self.output_path_label.setText(dir_path)
    
    def start_conversion(self):
        if self.file_list.count() == 0:
            QMessageBox.warning(self, "警告", "请先添加文件！")
            return
        
        if not self.output_dir:
            QMessageBox.warning(self, "警告", "请先选择输出目录！")
            return
        
        file_list = []
        for i in range(self.file_list.count()):
            file_list.append(self.file_list.item(i).data(Qt.UserRole))
        
        format_type = self.format_combo.currentText()
        pages = self.pages_spin.value()
        if pages == 999:
            pages = "all"
        
        # 根据清晰度选择设置DPI值
        clarity = self.dpi_combo.currentText()
        if clarity == '低':
            dpi = 96  # 适合屏幕显示
        elif clarity == '中':
            dpi = 200  # 适合一般打印
        else:  # 高
            dpi = 300  # 适合高质量印刷
        
        self.convert_btn.setEnabled(False)
        self.progress_bar.setValue(0)
        self.log_text.clear()
        
        self.converter_thread = ConverterThread(file_list, self.output_dir, format_type, pages, dpi)
        self.converter_thread.progress_updated.connect(self.update_progress)
        self.converter_thread.file_started.connect(self.log_file_started)
        self.converter_thread.conversion_finished.connect(self.conversion_complete)
        self.converter_thread.start()
    
    def update_progress(self, progress, info):
        self.progress_bar.setValue(progress)
        self.log_text.append(f"进度: {progress}% - {info}")
    
    def log_file_started(self, filename):
        self.log_text.append(f"开始处理: {filename}")
    
    def conversion_complete(self, success, message, completed, total):
        self.convert_btn.setEnabled(True)
        self.log_text.append(f"转换完成: {message}")
        
        if success:
            QMessageBox.information(self, "完成", message)
        else:
            QMessageBox.warning(self, "错误", message)
    
    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            
    def dropEvent(self, event: QDropEvent):
        files = [u.toLocalFile() for u in event.mimeData().urls()]
        for file_path in files:
            if os.path.isfile(file_path) and file_path.lower().endswith(('.pdf', '.docx')):
                if not self.is_file_in_list(file_path):
                    item = QListWidgetItem(file_path)
                    item.setData(Qt.UserRole, file_path)
                    self.file_list.addItem(item)
            elif os.path.isdir(file_path):
                # 如果是文件夹，递归添加所有支持的文件
                for root, dirs, files in os.walk(file_path):
                    for file in files:
                        if file.lower().endswith(('.pdf', '.docx')):
                            full_path = os.path.join(root, file)
                            if not self.is_file_in_list(full_path):
                                item = QListWidgetItem(full_path)
                                item.setData(Qt.UserRole, full_path)
                                self.file_list.addItem(item)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    converter = DocumentConverter()
    converter.show()
    sys.exit(app.exec_())